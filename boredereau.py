import pandas as pd
from docx import Document
from docx.shared import Pt

def format_lines(row):
    """Return a list of lines (strings) for one entry, given a dataframe row."""
    client = str(row.get("Client", "")).strip()
    commodity = str(row.get("Marchandise", "")).strip()
    nb_colis = row.get("nombre colis", "")
    tonnage = row.get("Poids brute", "")
    
    # Convert to string, maybe formatting decimals
    nb_str = str(nb_colis)
    tonnage_str = str(tonnage)
    
    lines = []
    lines.append("=*=*=*=*=*=*=*=*=*=*=*=*=*=*=*=*=*=")
    lines.append(f"Receiver : {client}    Comodity : {commodity}")
    lines.append(f"Manifested Quantity:  {nb_str}    COILS    Tonnage : {tonnage_str}  Mt")
    lines.append("    Received    Coils Packaging damaged on board.")
    lines.append(f"Total Received    {nb_str}    Coils.")
    lines.append("The Quantity Will Be confirmed after delivery Cargo.")
    lines.append("=*=*=*=*=*=*=*=*=*=*=*=*=*=*=*=*=*=")
    return lines

def excel_to_docx(input_excel, sheet_name=None, output_docx="output.docx"):
    # Load Excel
    df = pd.read_excel(input_excel, sheet_name=sheet_name)
    
    # Create Word document
    doc = Document()
    
    # Optionally set a default font or style
    # For example, set normal style font size:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    
    # For each row, add the entry
    for idx, row in df.iterrows():
        lines = format_lines(row)
        for line in lines:
            # Add a paragraph for each line
            p = doc.add_paragraph(line)
            # Optionally enforce formatting (e.g. no extra spacing)
            # p.paragraph_format.space_after = Pt(0)
        # Add a blank paragraph / line between entries
        doc.add_paragraph("")  
    
    # Save the document
    doc.save(output_docx)
    print(f"Saved {output_docx}")

if __name__ == "__main__":
    # Example usage — change filenames as needed
    excel_to_docx("Book1.xlsx", sheet_name=0, output_docx="entries.docx")
