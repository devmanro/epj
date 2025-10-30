from docx import Document
import pandas as pd

def replace_placeholders_in_paragraph(paragraph, replacements):
    """
    For a given paragraph, if its full text contains any placeholder key,
    replace it with the target value, reconstructing runs.
    """
    full_text = paragraph.text
    for placeholder, new_value in replacements.items():
        if placeholder in full_text:
            # Compute the replaced text
            replaced_text = full_text.replace(placeholder, new_value)
            # Remove all existing runs
            # (Be careful: modifying runs while iterating is tricky; do reversed)
            for run in reversed(paragraph.runs):
                r = run._element
                r.getparent().remove(r)
            # Add a single new run with the replaced text
            paragraph.add_run(replaced_text)
            break  # we assume one replacement per paragraph; adjust if multiple
    # Note: this loses run-level styling (bold, italic) inside replaced parts.
    # If you need to preserve styling around the fields, a more granular approach is needed.

def replace_in_docx_template(template_path, output_path, replacements):
    """
    Open an existing docx at template_path, replace all placeholder keys
    per the `replacements` dict, and save to output_path.
    """
    doc = Document(template_path)
    # Replace in normal paragraphs
    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, replacements)
    # Also replace in tables (cells) if template uses tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_in_paragraph(para, replacements)
    # Save the modified document
    doc.save(output_path)

def fill_from_excel_using_template(template_path, excel_path, output_prefix="filled"):
    df = pd.read_excel(excel_path, engine="openpyxl")
    for idx, row in df.iterrows():
        # Build the placeholder → actual value mapping
        replacements = {
            "Receiver :": f"Receiver : {row.get('Client', '')}",
            "commodity :": f"commodity : {row.get('Marchandise', '')}",
            "Manifested Quantity :": f"Manifested Quantity: {row.get('nombre colis', '')}",
            "tonnage :": f"tonnage: {row.get('Poids brute', '')}",
            # Add more keys if your template uses different names
        }
        outname = f"{output_prefix}_{idx+1}.docx"
        replace_in_docx_template(template_path, outname, replacements)
        print(f"Generated {outname}")

if __name__ == "__main__":
    fill_from_excel_using_template("src.docx", "Book1.xlsx", output_prefix="output")
