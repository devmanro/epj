import pandas as pd
from docx import Document
from docx.shared import Pt, Inches , Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def format_entry_docx(doc, row):
    client = str(row.get("client", "")).strip()
    commodity = str(row.get("type", "")).strip() or "Units + Package"
    nb_colis = row.get("qte", "") or 00
    tonnage = row.get("poids", "") or 00
    rec_qty = row.get("rec_qty", "") or 00

    # manifest_qty_str = str(nb_colis)
    tonnage_str = str(tonnage)
    manifest_qty_str = f"{int(nb_colis):02d}"
    rec_str = f"{int(rec_qty):02d}" 
    damaged_str=str("00")

    # create table with 2 columns for labels/values, style similar to your image
    table = doc.add_table(rows=5, cols=2)
    
    table.autofit = True
    # optional: fix widths
    # widths = (Inches(3), Inches(3))
    # table.columns[0].width = Inches(1)  # adjust as needed
    # table.columns[1].width = Inches(1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # for row_cells in table.rows:
    #     for idx, width in enumerate(widths):
    #         row_cells.cells[idx].width = width

    # for i, row in enumerate(table.rows):
    #     row.cells[0].width = Cm(12)
    #     row.cells[1].width = Cm(5)
    # Row 0: Receiver / Commodity
    row0 = table.rows[0].cells
    row0[0].width = Cm(9)
    row0[0].paragraphs[0].add_run("Receiver : ").bold = True
    row0[0].paragraphs[0].add_run(client)
    row0[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row0[1].width = Cm(9)
    row0[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    run=row0[1].paragraphs[0].add_run("Commodity : ")
    run1=row0[1].paragraphs[0].add_run(commodity)
    run.bold=True
    run.font.name = "Agency FB"
    run1.font.name = "Agency FB"


    # Row 1: Manifested Quantity / Tonnage
    row1 = table.rows[1].cells 
    row1[0].width = Cm(12)
    row1[0].paragraphs[0].add_run("Manifested Quantity : ").bold = True
    row1[0].paragraphs[0].add_run(f"{manifest_qty_str} UNIT + PACKAGE")
    row1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row1[1].width = Cm(5)
    row1[1].paragraphs[0].add_run("Tonnage : ").bold = True
    row1[1].paragraphs[0].add_run(f"{tonnage_str} Mt")
    row1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Row 2: Received / (right empty)
    row2 = table.rows[2].cells
    row2[0].width = Cm(30)
    row2[0].paragraphs[0].add_run("Received:   ").bold = True
    row2[0].paragraphs[0].add_run(f"    {damaged_str} Packaging damaged on board" ) 
    row2[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    # right cell left blank or you could merge if you like
    # row2[1].width = Cm(12)
    # row2[1].paragraphs[0].add_run(f"{damaged_str} Packaging damaged on board" ) 
    # row2[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Row 3: Total Received / (commodity)
    row3 = table.rows[3].cells
    row3[0].width = Cm(12)
    row3[0].paragraphs[0].add_run("Total Received:  ").bold = True
    row3[0].paragraphs[0].add_run(f"  {rec_str}")
    row3[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    # row3[0].paragraphs[0].add_run("")  # you could add value if needed

    # row3[1].width = Cm(5)
    # row3[1].paragraphs[0].add_run(rec_str)
    # row3[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Row 4: Final line (spanning both columns)
    # simplest: write in left cell and span visually
    row4 = table.rows[4].cells
    row4[0].width = Cm(25)
    full = row4[0].paragraphs[0].add_run(
        "The Quantity Will Be confirmed after delivery Cargo."
    )
    full.bold = True
    # optionally merge cells:
    # table.rows[4].cells[0]._tc.merge(table.rows[4].cells[1]._tc)

    # blank line after table
    doc.add_paragraph()

def excel_to_docx_custom(input_excel, sheet_name=None, template_path=None, output_docx="output.docx"):
    df = pd.read_excel(input_excel, sheet_name=sheet_name, engine="openpyxl")
    doc = Document(template_path) if template_path else Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri (Corps)"
    font.size = Pt(12)

    for idx, row in df.iterrows():
        format_entry_docx(doc, row)

    doc.save(output_docx)
    print(f"Saved {output_docx}")

if __name__ == "__main__":
    excel_to_docx_custom("Book1.xlsx", sheet_name=0, template_path="template.docx", output_docx="entries.docx")
