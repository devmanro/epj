import pandas as pd
from docx import Document
from docx.shared import Pt, Inches , Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.table import Table
from docx.document import Document as DocType # Type hinting for clarity

# --- New Helper Function for Space Calculation ---

def calculate_char_width(doc: DocType, table: Table, font_name="Courier New", font_size_pt=10):
    """
    Estimates the required character width for each column based on the
    physical width of the column in the Word document.
    """
    column_widths_chars = []
    
    # This is an approximation. We assume 1 point (Pt) is roughly 0.75 pixels, 
    # and a monospaced character is roughly 6-7 points wide at 10pt size.
    # We use a simple divisor to convert width (in EMUs) to character count.
    # 914400 EMUs = 1 Inch. 
    # A 10pt character is roughly ~70000 EMUs wide.
    EMUS_PER_CHAR = 70000 
    
    # Get the columns object from the table
    columns = table.columns 
    
    for col in columns:
        # Col width is in EMUs (English Metric Units)
        width_emus = col.width 
        
        # Calculate approximate character count
        char_count = int(width_emus / EMUS_PER_CHAR)
        
        # Ensure a minimum size to hold small values
        column_widths_chars.append(max(char_count, 10))
        
    return column_widths_chars

# --- Modified Conversion Function ---

def convert_and_delete_table(document: DocType, table: Table):
    """
    Converts a single, specified table into text separated by spaces to preserve
    column structure (requires a monospaced font), then deletes the table.
    """
    # 1. Calculate the target character width for each column
    column_widths = calculate_char_width(document, table)
    
    # --- Extract and Insert Content with Padding ---
    for row in table.rows:
        row_text_padded = []
        
        for i, cell in enumerate(row.cells):
            # Extract text (concatenating multi-paragraph cell content)
            cell_content = '\n'.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            
            # Get the target width for this column
            target_width = column_widths[i]
            
            # Pad the content with spaces to match the target width
            # We use an f-string left alignment and truncation for padding
            # '{:<{width}}' means left-align (<) and pad to 'width' size
            padded_content = f'{cell_content:<{target_width}}'
            
            # Truncate if content is too long to prevent alignment breaking
            if len(padded_content) > target_width:
                 padded_content = padded_content[:target_width] 
            
            row_text_padded.append(padded_content)
        
        # Join the padded cell contents without any separator (the spaces are the separator)
        final_row_text = "".join(row_text_padded)
        
        # Insert the new paragraph
        if final_row_text.strip():
            p = document.add_paragraph(final_row_text)
            
            # Set the new paragraph to a monospaced font to maintain column alignment
            run = p.runs[0]
            run.font.name = "Courier New" # MUST be monospaced
            run.font.size = Pt(10) # Set size for consistent character scaling
            
    # Optional: Add an extra line break after the converted table content
    document.add_paragraph() 
    
    # --- Delete the Table via XML manipulation ---
    tbl = table._element
    parent = tbl.getparent()
    if parent is not None:
        parent.remove(tbl)
    else:
        print(f"Warning: Could not find parent of table for deletion.")


# --- Modified Original Function to Return the Table ---

def format_entry_docx(doc: DocType, row: pd.Series) -> Table:
    """
    Creates a formatted table entry in the document and returns the created Table object.
    (Content generation logic remains the same)
    """
    client = str(row.get("client", "")).strip()
    commodity = str(row.get("type", "")).strip() or "Units + Package"
    nb_colis = row.get("qte") or 00
    tonnage = row.get("poids") or 0.0
    rec_qty = row.get("rec_qty") or 00
    
    tonnage_str = f"{tonnage:.2f}".lstrip("0") if tonnage < 1 and tonnage > 0 else f"{tonnage:.2f}"
    tonnage_str = f"{tonnage_str}"

    manifest_qty_str = f"{int(nb_colis):02d}"
    rec_str = f"{int(rec_qty):02d}"
    damaged_str=str("00")

    # create table with 2 columns for labels/values
    table = doc.add_table(rows=5, cols=2) 
    
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: Receiver / Commodity
    row0 = table.rows[0].cells
    row0[0].width = Cm(9)
    p00 = row0[0].paragraphs[0]
    p00.add_run("Receiver : ").bold = True
    p00.add_run(client)
    p00.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    row0[1].width = Cm(9)
    p01 = row0[1].paragraphs[0]
    p01.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_comm = p01.add_run("Commodity : ")
    run1_comm = p01.add_run(commodity)
    run_comm.bold=True
    run_comm.font.name = "Agency FB"
    run1_comm.font.name = "Agency FB"


    # Row 1: Manifested Quantity / Tonnage
    row1 = table.rows[1].cells 
    p10 = row1[0].paragraphs[0]
    p10.add_run("Manifested Quantity : ").bold = True
    p10.add_run(f"{manifest_qty_str} UNIT + PACKAGE")
    p10.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p11 = row1[1].paragraphs[0]
    p11.add_run("Tonnage : ").bold = True
    p11.add_run(f"{tonnage_str} Mt")
    p11.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Row 2: Received / (right empty)
    row2 = table.rows[2].cells
    p20 = row2[0].paragraphs[0]
    p20.add_run("Received:    ").bold = True
    p20.add_run(f"    {damaged_str} Packaging damaged on board" ) 
    p20.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Row 3: Total Received / (commodity)
    row3 = table.rows[3].cells
    p30 = row3[0].paragraphs[0]
    p30.add_run("Total Received:  ").bold = True
    p30.add_run(f"  {rec_str}")
    p30.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Row 4: Final line (spanning both columns)
    row4 = table.rows[4].cells
    p40 = row4[0].paragraphs[0]
    full = p40.add_run(
        "The Quantity Will Be confirmed after delivery Cargo."
    )
    full.bold = True
    
    return table # Return the table object


def excel_to_docx_custom(input_excel, sheet_name=None, template_path=None, output_docx="output.docx", convert_tables=True):
    # Load data and document
    df = pd.read_excel(input_excel, sheet_name=sheet_name, engine="openpyxl")
    doc = Document(template_path) if template_path else Document()

    # Set base style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri (Corps)"
    font.size = Pt(12)

    # 1. Loop and process each entry
    for idx, row in df.iterrows():
        # Step A: Create the table and get the object reference
        new_table = format_entry_docx(doc, row)
        
        # Step B: Immediately convert the table to space-padded text and delete the table
        if convert_tables:
            convert_and_delete_table(doc, new_table)

    # 2. Save the document
    doc.save(output_docx)
    print(f"Saved {output_docx}")

if __name__ == "__main__":
    excel_to_docx_custom(
        input_excel="Book1.xlsx", 
        sheet_name=0, 
        template_path="template.docx", 
        output_docx="entries_space_padded.docx",
        convert_tables=True
    )